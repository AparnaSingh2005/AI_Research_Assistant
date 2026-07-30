from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet


class Exporter:

    # -----------------------------------------
    # TXT
    # -----------------------------------------

    def export_txt(self, text, filename):

        with open(filename, "w", encoding="utf-8") as file:
            file.write(text)

        return filename

    # -----------------------------------------
    # DOCX
    # -----------------------------------------

    def export_docx(self, text, filename):

        document = Document()

        document.add_heading(
            "AI Research Paper Assistant",
            level=1
        )

        document.add_paragraph(text)

        document.save(filename)

        return filename

    # -----------------------------------------
    # PDF
    # -----------------------------------------

    def export_pdf(self, text, filename):

        pdf = SimpleDocTemplate(filename)

        styles = getSampleStyleSheet()

        story = []

        for line in text.split("\n"):

            story.append(
                Paragraph(
                    line,
                    styles["BodyText"]
                )
            )

        pdf.build(story)

        return filename